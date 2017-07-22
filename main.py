import openpyxl
from xml.etree.cElementTree import XML
from docx import Document
import zipfile

#Open document
def openWordDocument(filename):
    return Document(filename)

#Open work books
def openWorkbook(filename):
    return openpyxl.load_workbook(filename)

#Retrieve Sheet names
def getWorksheets(workbook):
    return workbook.get_sheet_names()

# Scan worksheet
def scanMenteeWorksheet(collection, worksheet):
    for row in range(2, worksheet.max_row + 1):
        person = {    
            'name' : worksheet['B' + str(row)].value,
            'email' : worksheet['D' + str(row)].value,
            'phone_number' : worksheet['AG' + str(row)].value,
            'description' : worksheet['K' + str(row)].value,
            'study_habits' : worksheet['S' + str(row)].value,
            'mentor' : worksheet['C' + str(row)].value,
            'hometown' : worksheet['F' + str(row)].value,
            'undergrad' : worksheet['L' + str(row)].value,
            'major' : worksheet['M' + str(row)].value,
            'music' : worksheet['Y' + str(row)].value,
            'movie' : worksheet['X' + str(row)].value,
            'hobbies' : worksheet['V' + str(row)].value
        }
        collection.append(person)

def scanMentorWorksheet(collection, worksheet):
    for row in range(2, worksheet.max_row + 1):
        person = {
            'name' : worksheet['B' + str(row)].value,
            'email' : worksheet['D' + str(row)].value,
            'description' : worksheet['L' + str(row)].value,
            'study_habits' : worksheet['T' + str(row)].value,
            'mentee' : worksheet['C' + str(row)].value,
            'hometown' : worksheet['G' + str(row)].value,
            'undergrad' : worksheet['M' + str(row)].value,
            'major' : worksheet['N' + str(row)].value,
            'music' : worksheet['Z' + str(row)].value,
            'movie' : worksheet['Y' + str(row)].value,
            'hobbies' : worksheet['W' + str(row)].value
        }
        collection.append(person)

# Map mentor to mentee
def mapMentorMentee(collection, mentors, mentees):
    for mentor in mentors:
        for mentee in mentees:
            if mentor['name'] == mentee['mentor'] and mentor['mentee'] == mentee['name']:
                collection.append((mentor, mentee))

# Validate Mentee data.
def validateMentees(collection):
    for mentee in collection:
        if (
            mentee['email'] == '' or
            mentee['phone_number'] == '' or
            mentee['description'] == '' or
            mentee['study_habits'] == '' or
            mentee['mentor'] == '' or
            mentee['hometown'] == '' or
            mentee['undergrad'] == '' or
            mentee['major'] == '' or
            mentee['music'] == '' or
            mentee['movie'] == '' or
            mentee['hobbies']
        ):
            print(mentee['name'] + ' has some empty fields.')

# Validate Mentore Data
def validateMentors(collection):
    for mentor in collection:
        if (
            mentor['email'] == '' or
            mentor['description'] == '' or
            mentor['study_habits'] == '' or
            mentor['mentee'] == '' or
            mentor['hometown'] == '' or
            mentor['undergrad'] == '' or
            mentor['major'] == '' or
            mentor['music'] == '' or
            mentor['movie'] == '' or
            mentor['hobbies']
        ):
            print(mentor['name'] + ' has some empty fields.')

def generateWordDoc(matchCollection):
    MENTOR = 0
    MENTEE = 1
    old_file_name = 'example bio card.docx'
    for match in matchCollection:
        filename = match[MENTOR]['name'] + '-' + match[MENTEE]['name'] + '.docx'
        filename2 = match[MENTEE]['name'] + '-' + match[MENTOR]['name'] + '.docx'
        _ = Document()
        zin = zipfile.ZipFile (old_file_name, 'r')
        zout = zipfile.ZipFile (filename, 'w')
        zout2 = zipfile.ZipFile (filename2, 'w')

        for item in zin.infolist():
            buffer = zin.read(item.filename)
            if item.filename == 'word/document.xml':
                res = buffer.decode('utf-8')
                res = res.replace('Home town:', 'Home town: ' + match[MENTOR]['hometown'])
                res = res.replace('Undergrad school &amp; major:', 'Undergrad school &amp; major: ' + match[MENTOR]['undergrad']+', ' + match[MENTOR]['major'])
                res = res.replace('Favorite Music &amp; Movie:', 'Favorite Music &amp; Movie: ' + match[MENTOR]['music'] + ', ' + match[MENTOR]['movie'])
                res = res.replace('RECIPIENT NAME', match[MENTOR]['name'].upper())
                res = res.replace('Email', match[MENTOR]['email'])
                res = res.replace('phone', '') # Mentors dont have phone numbers yet
                res = res.replace('Persistent, Detail-Oriented, Honest', match[MENTOR]['description'])
                res = res.replace('Take Starbucks and endless snacks as needed for studying.', match[MENTOR]['study_habits'])
                res = res.replace('Hobbies:', 'Hobbies: ' + match[MENTOR]['hobbies'])
                res = res.replace('Sender Name', match[MENTEE]['name'])
                buffer = res.encode('utf-8')
            zout.writestr(item, buffer)
        zout.close()
        zin.close()
        
        zin = zipfile.ZipFile (old_file_name, 'r')
        for item in zin.infolist():
            buffer = zin.read(item.filename)
            if item.filename == 'word/document.xml':
                res = buffer.decode('utf-8')
                res = res.replace('Home town:', 'Home town: ' + match[MENTEE]['hometown'])
                res = res.replace('Undergrad school &amp; major:', 'Undergrad school &amp; major: ' + match[MENTEE]['undergrad']+', ' + match[MENTEE]['major'])
                res = res.replace('Favorite Music &amp; Movie:', 'Favorite Music &amp; Movie: ' + match[MENTEE]['music'] + ', ' + match[MENTEE]['movie'])
                res = res.replace('RECIPIENT NAME', match[MENTEE]['name'].upper())
                res = res.replace('email', match[MENTEE]['email'])
                res = res.replace('phone', str(match[MENTEE]['phone_number']))
                res = res.replace('Persistent, Detail-Oriented, Honest', match[MENTEE]['description'])
                res = res.replace('Take Starbucks and endless snacks as needed for studying.', match[MENTEE]['study_habits'])
                res = res.replace('Hobbies:', 'Hobbies: ' + match[MENTEE]['hobbies'])
                res = res.replace('Sender Name', match[MENTOR]['name'])
                buffer = res.encode('utf-8')
            zout2.writestr(item, buffer)
        zout2.close()
        zin.close()
        
# Main Function
def main():
    print('Opening documents...')
    biocard = openWordDocument('example bio card.docx')
    menteeWorkbook = openWorkbook('Mentee Survey (Responses).xlsx')
    mentorWorkbook = openWorkbook('Mentor Survey (Responses).xlsx')

    print('Retrieving worksheets...')
    menteeWorksheets = getWorksheets(menteeWorkbook)
    mentorWorksheets = getWorksheets(mentorWorkbook)

    menteeWorksheet = menteeWorkbook.get_sheet_by_name(menteeWorksheets[0])
    mentorWorksheet = mentorWorkbook.get_sheet_by_name(mentorWorksheets[0])

    print('Scanning rows...')
    #Find Mentor Mentee Pair
    mentees = [] 
    mentors = []
    match = []
    
    print('Scanning mentees...')
    scanMenteeWorksheet(mentees, menteeWorksheet)
    
    print('Scanning mentors...')
    scanMentorWorksheet(mentors, mentorWorksheet)
    
    print('Validating mentee/mentor answers...')
    validateMentees(mentees)
    validateMentors(mentors)
    print('Answers validated...')
    
    print('Mapping mentors to mentees...')
    mapMentorMentee(match, mentors, mentees)

    print(str(len(match)) + ' pairs have been created. Now generating documents for the matches...')
    generateWordDoc(match)

if __name__ == '__main__':
    main()